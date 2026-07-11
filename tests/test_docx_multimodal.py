"""test_docx_multimodal.py — M-4 회귀 가드.

extract_docx_multimodal 의 DOCX 표 추출이:
  (1) 병합셀(gridSpan/vMerge)을 보존(값 중복 전개 금지),
  (2) GFM(GitHub Flavored Markdown) 파이프 표를 직접 생성,
  (3) 길이 기준 청크 분할,
  (4) truncation(finish_reason=length) 감지 → .partial.md 저장
하는지 검증한다. LLM 호출은 ox.extract_text_prompt 를 stub 으로 교체.

실행:
    .venv/bin/python -m pytest tests/test_docx_multimodal.py -v
"""
from __future__ import annotations

import os
import sys
import tempfile
import unittest
from pathlib import Path
from unittest.mock import MagicMock, patch

from docx import Document

_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if _ROOT not in sys.path:
    sys.path.insert(0, _ROOT)

import extract_docx_multimodal as d  # noqa: E402


def _save_tmp(doc) -> Path:
    fd, p = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    doc.save(p)
    return Path(p)


# ──────────────────────────────────────────────────────────────────────────────
# (1) 병합셀 보존 — 가로(gridSpan)
# ──────────────────────────────────────────────────────────────────────────────

class TestHorizontalMergePreserved(unittest.TestCase):

    def test_hmerge_not_duplicated(self):
        """가로 병합 셀 값이 중복 전개되지 않고 1회 + 빈 패딩으로 열 구조 보존."""
        doc = Document()
        t = doc.add_table(rows=2, cols=3)
        t.cell(0, 0).text = "H1"
        t.cell(0, 1).text = "H2"
        t.cell(0, 2).text = "H3"
        a = t.cell(1, 0)
        b = t.cell(1, 1)
        a.merge(b)
        a.text = "MERGED"
        t.cell(1, 2).text = "val"
        p = _save_tmp(doc)
        try:
            blocks = d.extract_blocks_from_docx(p)
        finally:
            p.unlink()

        gfm = blocks[0]
        # 병합 값은 정확히 1회만 등장(중복 전개 X — `MERGED | MERGED` 금지).
        self.assertEqual(gfm.count("MERGED"), 1, gfm)
        # 데이터 행: 3열 유지(파이프 4개) — 병합 칸은 빈 셀로 채워짐.
        data_row = [ln for ln in gfm.splitlines() if "MERGED" in ln][0]
        self.assertEqual(data_row.count("|"), 4, data_row)
        # 빈 패딩 셀 존재(연속 파이프 사이 공백).
        self.assertIn("|  |", data_row, data_row)


# ──────────────────────────────────────────────────────────────────────────────
# (1) 병합셀 보존 — 세로(vMerge)
# ──────────────────────────────────────────────────────────────────────────────

class TestVerticalMergePreserved(unittest.TestCase):

    def test_vmerge_continuation_empty(self):
        """세로 병합 연속 셀은 빈 값(값 행마다 반복 전개 금지)."""
        doc = Document()
        t = doc.add_table(rows=3, cols=2)
        t.cell(0, 0).text = "A"
        t.cell(0, 1).text = "B"
        top = t.cell(1, 0)
        bot = t.cell(2, 0)
        top.merge(bot)
        top.text = "VMERGED"
        t.cell(1, 1).text = "x"
        t.cell(2, 1).text = "y"
        p = _save_tmp(doc)
        try:
            blocks = d.extract_blocks_from_docx(p)
        finally:
            p.unlink()

        gfm = blocks[0]
        # 세로 병합 값은 정확히 1회(origin 행에만) — 연속 행은 빈 셀.
        self.assertEqual(gfm.count("VMERGED"), 1, gfm)
        # 마지막 데이터 행(y 있는 행)은 첫 셀이 비어 있음.
        last = [ln for ln in gfm.splitlines() if "y" in ln][0]
        # `|  | y |` 형태 — 첫 셀 공백.
        self.assertTrue(last.lstrip().startswith("|  |"), last)


# ──────────────────────────────────────────────────────────────────────────────
# (2) GFM 직접 생성 — str(list) 덤프 금지
# ──────────────────────────────────────────────────────────────────────────────

class TestGfmDirect(unittest.TestCase):

    def test_table_is_gfm_not_python_list(self):
        doc = Document()
        t = doc.add_table(rows=2, cols=2)
        t.cell(0, 0).text = "Name"
        t.cell(0, 1).text = "Value"
        t.cell(1, 0).text = "R1"
        t.cell(1, 1).text = "10K"
        p = _save_tmp(doc)
        try:
            blocks = d.extract_blocks_from_docx(p)
        finally:
            p.unlink()
        gfm = blocks[0]
        # GFM 헤더 구분선 존재.
        self.assertIn("| --- |", gfm)
        self.assertIn("| Name | Value |", gfm)
        self.assertIn("| R1 | 10K |", gfm)
        # 파이썬 리스트 덤프 흔적 부재.
        self.assertNotIn("[TABLE]", gfm)
        self.assertNotIn("[['", gfm)

    def test_pipe_in_cell_escaped(self):
        """셀 내부 파이프는 이스케이프되어 표 구조를 깨지 않음."""
        doc = Document()
        t = doc.add_table(rows=1, cols=1)
        t.cell(0, 0).text = "a|b"
        p = _save_tmp(doc)
        try:
            blocks = d.extract_blocks_from_docx(p)
        finally:
            p.unlink()
        self.assertIn("a\\|b", blocks[0])


# ──────────────────────────────────────────────────────────────────────────────
# (3) 길이 기준 청크 분할
# ──────────────────────────────────────────────────────────────────────────────

class TestChunking(unittest.TestCase):

    def test_splits_on_block_boundary(self):
        chunks = d.chunk_blocks(["a" * 5000, "b" * 5000, "c" * 5000], max_chars=8000)
        self.assertEqual(len(chunks), 3)

    def test_groups_small_blocks(self):
        chunks = d.chunk_blocks(["a" * 1000, "b" * 1000, "c" * 1000], max_chars=8000)
        self.assertEqual(len(chunks), 1)

    def test_oversized_single_block_kept_whole(self):
        """단일 블록이 max_chars 초과해도 단독 청크로 보존(표 무결성)."""
        chunks = d.chunk_blocks(["z" * 20000], max_chars=8000)
        self.assertEqual(len(chunks), 1)
        self.assertEqual(len(chunks[0]), 20000)

    def test_empty(self):
        self.assertEqual(d.chunk_blocks([], 8000), [])


# ──────────────────────────────────────────────────────────────────────────────
# (4) truncation 감지 → .partial.md
# ──────────────────────────────────────────────────────────────────────────────

class TestTruncationDetection(unittest.TestCase):

    def setUp(self):
        self._tmpdir = tempfile.TemporaryDirectory()
        self._tmp = Path(self._tmpdir.name)
        # OUTPUT_DIR 을 임시로 교체.
        self._orig_out = d.OUTPUT_DIR
        d.OUTPUT_DIR = self._tmp

    def tearDown(self):
        d.OUTPUT_DIR = self._orig_out
        self._tmpdir.cleanup()

    def _make_docx(self) -> Path:
        doc = Document()
        doc.add_paragraph("Some content")
        t = doc.add_table(rows=1, cols=2)
        t.cell(0, 0).text = "k"
        t.cell(0, 1).text = "v"
        # 출력 파일명 안정화를 위해 stem 고정.
        p = self._tmp / "mydoc.docx"
        doc.save(p)
        return p

    def test_truncated_saves_partial(self):
        """LLM 응답에 TRUNCATED 마커 → .partial.md 로 저장(완성본 위장 금지)."""
        docx = self._make_docx()
        stub = MagicMock(return_value="## ok\n<!-- TRUNCATED: finish_reason=length -->")
        with patch.object(d.ox, "extract_text_prompt", stub), \
             patch.object(d.ox, "provider_label", lambda: "stub"), \
             patch.object(d, "time"):
            d.process_docx(docx)
        partial = self._tmp / "mydoc.partial.md"
        final = self._tmp / "mydoc.md"
        self.assertTrue(partial.exists(), ".partial.md 가 생성되어야 함")
        self.assertFalse(final.exists(), "완성본 .md 가 생성되면 안 됨")

    def test_clean_saves_final(self):
        """truncation 없으면 .md 로 저장."""
        docx = self._make_docx()
        stub = MagicMock(return_value="## ok\nclean content")
        with patch.object(d.ox, "extract_text_prompt", stub), \
             patch.object(d.ox, "provider_label", lambda: "stub"), \
             patch.object(d, "time"):
            d.process_docx(docx)
        final = self._tmp / "mydoc.md"
        partial = self._tmp / "mydoc.partial.md"
        self.assertTrue(final.exists(), ".md 가 생성되어야 함")
        self.assertFalse(partial.exists(), ".partial.md 가 생성되면 안 됨")

    def test_multi_chunk_joined_with_sep(self):
        """다청크 결과는 \\n\\n---\\n\\n 으로 결합."""
        docx = self._make_docx()
        # 청크가 2개 나오도록 DOCX_CHUNK_CHARS 를 낮춘다.
        outputs = iter(["chunk_a", "chunk_b"])
        stub = MagicMock(side_effect=lambda prompt: next(outputs))
        with patch.object(d.ox, "extract_text_prompt", stub), \
             patch.object(d.ox, "provider_label", lambda: "stub"), \
             patch.object(d, "DOCX_CHUNK_CHARS", 1), \
             patch.object(d, "time"):
            d.process_docx(docx)
        final = self._tmp / "mydoc.md"
        self.assertTrue(final.exists())
        content = final.read_text(encoding="utf-8")
        self.assertIn("\n\n---\n\n", content)


if __name__ == "__main__":
    unittest.main(verbosity=2)
