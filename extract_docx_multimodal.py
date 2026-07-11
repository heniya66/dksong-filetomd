import os
import sys
import time
from pathlib import Path
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

# 공통 추출 모듈(fmdw.ollama_extractor) import 경로 보장 — 워크스페이스 루트 추가.
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from fmdw import ollama_extractor as ox  # noqa: E402

# 추출 provider: 기본 ollama_cloud(로컬 게이트웨이, 키 불필요),
# EXTRACT_PROVIDER=gemini 로 기존 Gemini File API 경로 fallback.
# 본 스크립트는 docx 텍스트/표를 추출해 LLM으로 Markdown 정형화하는 비-멀티모달 경로.

# Directories
INPUT_DIR = Path("input/docx")
OUTPUT_DIR = Path("output/docx_md")
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# Word(OOXML) 네임스페이스 — gridSpan/vMerge 속성 조회용.
_W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

# M-4: 청크 분할 기준 문자 길이(추출 컨텍스트 블록의 누적 길이). 이 길이를 넘으면
#      블록 경계에서 다음 청크로 넘긴다. env DOCX_CHUNK_CHARS 로 override.
#      문서 전체를 무청크 단일 호출로 보내면 대형 DOCX 에서 응답이 토큰 한도로 잘려
#      (finish_reason=length) 표/값이 조용히 소실된다 → 길이 기준 분할로 완화.
try:
    DOCX_CHUNK_CHARS: int = int(os.getenv("DOCX_CHUNK_CHARS", "12000"))
except ValueError:
    DOCX_CHUNK_CHARS = 12000

# 청크 간 rate-limit 대기(초). env DOCX_RATE_DELAY 로 override.
try:
    DOCX_RATE_DELAY: float = float(os.getenv("DOCX_RATE_DELAY", "5"))
except ValueError:
    DOCX_RATE_DELAY = 5.0

# 청크 결합 구분자(extract_all_via_pdf 와 동일 계약).
CHUNK_SEP = "\n\n---\n\n"

# H-4 연계: LLM 응답이 토큰 한도로 잘리면 ox._ollama_vision 가 본문에 삽입하는 마커.
TRUNCATED_MARKER = "<!-- TRUNCATED"


def _iter_block_items(parent):
    """문서 본문을 등장 순서대로 (Paragraph | Table) 로 순회한다.

    기존 구현은 `doc.element.body` 를 돌며 매 블록마다 `doc.paragraphs`/`doc.tables`
    전체를 선형 탐색(O(n^2))했다. 여기서는 python-docx 객체로 직접 래핑한다.
    """
    body = parent.element.body
    for child in body.iterchildren():
        if child.tag == f"{_W_NS}p":
            yield Paragraph(child, parent)
        elif child.tag == f"{_W_NS}tbl":
            yield Table(child, parent)


def _grid_span(cell) -> int:
    """셀의 가로 병합 폭(w:gridSpan). 미지정이면 1."""
    tcPr = cell._tc.find(f"{_W_NS}tcPr")
    if tcPr is None:
        return 1
    grid = tcPr.find(f"{_W_NS}gridSpan")
    if grid is None:
        return 1
    try:
        return max(1, int(grid.get(f"{_W_NS}val", "1")))
    except (TypeError, ValueError):
        return 1


def _vmerge_state(cell) -> str:
    """세로 병합 상태: 'restart'(병합 시작/단일), 'continue'(병합 연속).

    w:vMerge 가 없으면 'restart'(병합 아님). val="restart" 도 시작. val 미지정/"continue"
    는 연속 셀(위 셀의 값에 종속) → 빈 값으로 출력해 값 중복 전개를 막는다.
    """
    tcPr = cell._tc.find(f"{_W_NS}tcPr")
    if tcPr is None:
        return "restart"
    vm = tcPr.find(f"{_W_NS}vMerge")
    if vm is None:
        return "restart"
    val = (vm.get(f"{_W_NS}val") or "continue").lower()
    return "restart" if val == "restart" else "continue"


def _escape_cell(text: str) -> str:
    """GFM 파이프 표 셀 이스케이프 — 파이프/개행 보존."""
    return text.replace("\\", "\\\\").replace("|", "\\|").replace("\n", "<br>").strip()


def table_to_gfm(table: Table) -> str:
    """python-docx Table → GFM 파이프 표 문자열(병합셀 보존).

    M-4 핵심:
      - 가로 병합(gridSpan): python-docx 의 `row.cells` 는 병합 셀을 **같은 값으로
        N회 중복** 반환한다. 여기서는 `tr/tc` 를 직접 순회하여 셀당 1회만 값을 쓰고,
        gridSpan 만큼 **빈 셀**로 폭을 채워 열 정렬과 구조를 보존한다(값 중복 제거).
      - 세로 병합(vMerge continue): 연속 셀은 빈 값으로 출력(원본은 위 origin 셀에만
        값 존재) → 같은 값을 행마다 반복 전개하지 않는다.
      - 열 수는 grid 폭(gridSpan 합)의 최댓값으로 정규화해 GFM 표 무결성 유지.
    """
    rows: list[list[str]] = []
    max_cols = 0
    for tr in table._tbl.tr_lst:
        cells: list[str] = []
        for tc in tr.tc_lst:
            from docx.table import _Cell

            cell = _Cell(tc, table)
            span = _grid_span(cell)
            vmerge = _vmerge_state(cell)
            value = "" if vmerge == "continue" else _escape_cell(cell.text)
            cells.append(value)
            # gridSpan>1 → 나머지 폭을 빈 셀로 채워 열 구조 보존(중복 값 전개 방지).
            for _ in range(span - 1):
                cells.append("")
        rows.append(cells)
        max_cols = max(max_cols, len(cells))

    if not rows or max_cols == 0:
        return ""

    # 열 수 정규화(짧은 행은 빈 셀로 패딩).
    for r in rows:
        if len(r) < max_cols:
            r.extend([""] * (max_cols - len(r)))

    lines: list[str] = []
    header = rows[0]
    lines.append("| " + " | ".join(header) + " |")
    lines.append("| " + " | ".join(["---"] * max_cols) + " |")
    for r in rows[1:]:
        lines.append("| " + " | ".join(r) + " |")
    return "\n".join(lines)


def extract_blocks_from_docx(docx_path) -> list[str]:
    """DOCX 본문을 등장 순서의 Markdown 블록 리스트로 추출한다.

    문단은 텍스트(빈 문단 제외), 표는 GFM 파이프 표(병합셀 보존)로 변환한다.
    반환 리스트는 길이 기준 청크 분할의 입력이 된다(블록 경계로만 분할).
    """
    doc = Document(docx_path)
    blocks: list[str] = []
    for item in _iter_block_items(doc):
        if isinstance(item, Paragraph):
            text = item.text.strip()
            if text:
                blocks.append(text)
        elif isinstance(item, Table):
            gfm = table_to_gfm(item)
            if gfm:
                blocks.append(gfm)
    return blocks


def chunk_blocks(blocks: list[str], max_chars: int) -> list[str]:
    """블록 리스트를 누적 길이 max_chars 기준으로 청크(블록 그룹)로 묶는다.

    블록 경계에서만 분할한다(표/문단을 중간에 자르지 않음). 단일 블록이 max_chars 를
    넘어도 그 블록은 단독 청크로 보존한다(표 무결성 우선 — LLM 측 truncation 마커로 감지).
    """
    if not blocks:
        return []
    chunks: list[str] = []
    cur: list[str] = []
    cur_len = 0
    for b in blocks:
        b_len = len(b)
        if cur and cur_len + b_len > max_chars:
            chunks.append("\n\n".join(cur))
            cur = []
            cur_len = 0
        cur.append(b)
        cur_len += b_len + 2  # "\n\n" 결합 오버헤드 근사
    if cur:
        chunks.append("\n\n".join(cur))
    return chunks


def _build_prompt(raw_block: str) -> str:
    """청크(이미 GFM 표를 포함한 raw 블록)를 정제 Markdown 으로 변환하는 프롬프트.

    표는 이미 GFM 로 추출되었으므로 LLM 은 표 구조를 **보존**하고 텍스트만 정형화한다.
    """
    return (
        "The following is content extracted from a DOCX file. Paragraphs are plain "
        "text and tables are ALREADY in GFM (GitHub Flavored Markdown) pipe-table "
        "format with merged cells preserved (merged regions appear as empty cells). "
        "Convert it into professional, high-quality Markdown.\n"
        "Guidelines:\n"
        "- Use hierarchical Markdown headings (#, ##, ###) based on the context.\n"
        "- PRESERVE the existing GFM pipe tables EXACTLY — do not collapse, re-expand, "
        "or duplicate merged (empty) cells; keep the column structure intact.\n"
        "- Clean up any transcription artifacts or repeated text in prose only.\n"
        "- Maintain all technical values, units, and signal names exactly as given.\n"
        "- Output ONLY the final Markdown content.\n\n"
        "RAW DATA:\n"
        "----------\n"
        f"{raw_block}"
    )


def process_docx(docx_path):
    print(f"\n[*] Processing: {docx_path.name}")

    try:
        # 1. Extract local content (blocks, in document order; tables → GFM).
        print(f"[*] Extracting text/tables from {docx_path.name}...")
        blocks = extract_blocks_from_docx(docx_path)

        if not blocks:
            print(f"[!] No content found in {docx_path.name}")
            return

        # 2. 길이 기준 청크 분할(무청크 단일 호출 → truncation 위험 완화).
        chunks = chunk_blocks(blocks, DOCX_CHUNK_CHARS)
        print(f"[*] {len(blocks)} block(s) → {len(chunks)} chunk(s) "
              f"(max {DOCX_CHUNK_CHARS} chars)", flush=True)

        # 3. 청크별 LLM 정형화 (provider 추상화 경유).
        md_parts: list[str] = []
        truncated = False
        for i, chunk in enumerate(chunks, start=1):
            print(f"    - Chunk {i}/{len(chunks)} via {ox.provider_label()}...",
                  flush=True)
            part = ox.extract_text_prompt(_build_prompt(chunk))
            md_parts.append(part)
            # M-4: truncation 감지 — H-4 가 본문에 삽입한 마커를 확인.
            if part and TRUNCATED_MARKER in part:
                truncated = True
                print(f"    [!] Chunk {i} truncated (finish_reason=length 의심)",
                      flush=True)
            if i < len(chunks) and DOCX_RATE_DELAY > 0:
                time.sleep(DOCX_RATE_DELAY)

        md_text = CHUNK_SEP.join(md_parts)

        # 4. Save Output — truncation 발생 시 완성본(.md)으로 위장하지 않는다.
        if truncated:
            output_path = OUTPUT_DIR / f"{docx_path.stem}.partial.md"
            print(f"[!] Truncation detected — saving partial: {output_path}", flush=True)
        else:
            output_path = OUTPUT_DIR / f"{docx_path.stem}.md"
        output_path.write_text(md_text, encoding="utf-8")
        print(f"[+] Successfully converted and saved to {output_path}")

    except Exception as e:
        print(f"[!] Error processing {docx_path.name}: {e}")


def main():
    if not INPUT_DIR.exists():
        print(f"Error: {INPUT_DIR} not found.")
        return

    docx_files = sorted(list(INPUT_DIR.glob("*.docx")))
    if not docx_files:
        print(f"No DOCX files found in {INPUT_DIR}")
        return

    print(f"Found {len(docx_files)} DOCX file(s) to process.")
    for docx in docx_files:
        process_docx(docx)
        time.sleep(5)  # Delay for rate limiting

if __name__ == "__main__":
    main()
